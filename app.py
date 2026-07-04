# -*- coding: utf-8 -*-
import os
import re
import tempfile
import subprocess
import zipfile
import sys
import base64
import io
from urllib.parse import urlparse
from pathlib import Path

try:
    from PIL import Image
except Exception:
    Image = None
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
import yt_dlp
from yt_dlp.utils import DownloadError

APP_NAME = "Transcreve Fácil"
APP_VERSION = "v18.16 - fragmentador local robusto"
ASSET_DIR = Path(__file__).parent / "assets"
LOGO_FULL = ASSET_DIR / "logo_full.png"
LOGO_ICON = ASSET_DIR / "logo_icon.png"
ALLOWED_DOMAIN = "@tre-ba.jus.br"
DEFAULT_USER = "vmsoares@tre-ba.jus.br"
DEFAULT_PASSWORD = "transcreve123"


def inject_css():
    st.markdown("""
    <style>
    :root {
        --tf-navy: #0b2f6b;
        --tf-blue: #1264f4;
        --tf-cyan: #00b8c8;
        --tf-teal: #12b8b8;
        --tf-orange: #ff7a1a;
        --tf-bg: #f4f9ff;
        --tf-border: #d7e8ff;
        --tf-text: #14213d;
        --tf-muted: #667799;
    }
    .stApp {
        background:
            radial-gradient(circle at top left, rgba(18, 184, 184, .13), transparent 34%),
            radial-gradient(circle at top right, rgba(18, 100, 244, .12), transparent 32%),
            linear-gradient(135deg, #f9fcff 0%, #eef7ff 100%);
        color: var(--tf-text);
    }
    [data-testid="stHeader"] {
        background: rgba(255, 255, 255, .78) !important;
        backdrop-filter: blur(10px);
        border-bottom: 1px solid rgba(215, 232, 255, .85);
    }
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff 0%, #f2f8ff 100%) !important;
        border-right: 1px solid var(--tf-border);
        box-shadow: 8px 0 28px rgba(11, 47, 107, .06);
    }
    section[data-testid="stSidebar"] > div {
        padding-top: 1.25rem;
    }
    .block-container {
        padding-top: 2.2rem !important;
        max-width: 1320px;
    }
    h1, h2, h3, h4 {
        color: #10264b;
        letter-spacing: -.02em;
    }
    .tf-topbar {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 1rem;
        background: rgba(255,255,255,.86);
        border: 1px solid var(--tf-border);
        border-radius: 24px;
        padding: .85rem 1rem;
        margin-bottom: 1.1rem;
        box-shadow: 0 16px 38px rgba(16, 38, 75, .07);
    }
    .tf-search {
        flex: 1;
        background: #fff;
        border: 1px solid #dceafd;
        border-radius: 16px;
        padding: .75rem 1rem;
        color: #7b8ca8;
        font-weight: 600;
    }
    .tf-user {
        display: flex;
        align-items: center;
        gap: .7rem;
        white-space: nowrap;
        font-weight: 800;
        color: var(--tf-navy);
    }
    .tf-avatar {
        height: 42px;
        width: 42px;
        border-radius: 50%;
        display: inline-flex;
        align-items: center;
        justify-content: center;
        color: white;
        background: linear-gradient(135deg, var(--tf-blue), var(--tf-cyan));
        box-shadow: 0 12px 26px rgba(18, 100, 244, .22);
    }
    .tf-hero {
        padding: 2rem 2.1rem;
        border-radius: 28px;
        background:
            linear-gradient(135deg, rgba(255,255,255,.98) 0%, rgba(237,248,255,.98) 100%);
        border: 1px solid #cfe4ff;
        box-shadow: 0 24px 60px rgba(2, 55, 110, .09);
        margin-bottom: 1.2rem;
        position: relative;
        overflow: hidden;
    }
    .tf-hero:after {
        content: "";
        position: absolute;
        width: 180px;
        height: 180px;
        right: -70px;
        top: -60px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(0,184,200,.18), transparent 68%);
    }
    .tf-hero h1 {
        font-size: 2.35rem;
        line-height: 1.1;
        margin: 0 0 .65rem 0;
        color: #10264b;
    }
    .tf-hero p {
        color: var(--tf-muted);
        font-size: 1.02rem;
        margin-bottom: .65rem;
    }
    .tf-badge {
        display: inline-flex;
        align-items: center;
        gap: .35rem;
        padding: .48rem .78rem;
        border-radius: 999px;
        margin-right: .42rem;
        margin-top: .38rem;
        background: #e6fffb;
        color: #036672;
        border: 1px solid #bdeff0;
        font-weight: 800;
        font-size: .9rem;
    }
    .tf-badge.blue { background: #eef5ff; color: #1155cc; border-color: #cfe0ff; }
    .tf-badge.orange { background: #fff3e9; color: #b84e00; border-color: #ffd5b1; }
    .tf-card {
        padding: 1.25rem;
        border-radius: 24px;
        background: rgba(255,255,255,.94);
        border: 1px solid var(--tf-border);
        box-shadow: 0 18px 42px rgba(16, 38, 75, .07);
        margin-bottom: 1rem;
    }
    .tf-card h3 { margin-top: 0; }
    .tf-icon {
        width: 54px;
        height: 54px;
        display: inline-flex;
        justify-content: center;
        align-items: center;
        border-radius: 16px;
        font-size: 1.6rem;
        margin-bottom: .7rem;
        box-shadow: 0 12px 24px rgba(16, 38, 75, .12);
    }
    .tf-icon.teal { background: linear-gradient(135deg, #15c7bd, #00a9d6); }
    .tf-icon.blue { background: linear-gradient(135deg, #246bfe, #2fa7ff); }
    .tf-icon.orange { background: linear-gradient(135deg, #ff7a1a, #ffbd4a); }
    .tf-uploadbox {
        border: 2px dashed #a9cdfc;
        border-radius: 24px;
        padding: 2rem;
        text-align: center;
        background: linear-gradient(180deg, rgba(255,255,255,.86), rgba(240,249,255,.7));
        margin: 1rem 0;
    }
    div.stButton > button, div.stDownloadButton > button, div[data-testid="stFormSubmitButton"] button {
        border-radius: 16px !important;
        border: 1px solid #cfe4ff !important;
        box-shadow: 0 10px 24px rgba(20, 82, 180, .10);
        font-weight: 800 !important;
        min-height: 2.8rem;
        transition: all .16s ease-in-out;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover, div[data-testid="stFormSubmitButton"] button:hover {
        border-color: #02aeca !important;
        color: #064e6b !important;
        transform: translateY(-1px);
        box-shadow: 0 14px 30px rgba(20, 82, 180, .14);
    }
    button[kind="primary"] {
        background: linear-gradient(135deg, var(--tf-blue), var(--tf-cyan)) !important;
        border: 0 !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: .45rem;
        background: rgba(255,255,255,.68);
        padding: .35rem;
        border-radius: 18px;
        border: 1px solid var(--tf-border);
        margin-bottom: .75rem;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 14px;
        padding: .6rem .9rem;
        font-weight: 800;
    }
    div[data-testid="stMetric"] {
        background: rgba(255,255,255,.78);
        border: 1px solid #dfebfb;
        border-radius: 18px;
        padding: .85rem;
    }
    .stAlert { border-radius: 16px; }
    textarea, input, .stSelectbox div[data-baseweb="select"] > div {
        border-radius: 14px !important;
    }

    .tf-brand-row {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin-bottom: 1rem;
    }
    .tf-brand-row img {
        width: 88px;
        height: 88px;
        object-fit: contain;
        border-radius: 24px;
        filter: drop-shadow(0 12px 24px rgba(18,100,244,.18));
    }
    .tf-brand-title {
        line-height: 1.03;
    }
    .tf-brand-title .one {
        font-size: 2rem;
        font-weight: 950;
        color: var(--tf-navy);
        letter-spacing: -.035em;
    }
    .tf-brand-title .two {
        font-size: 2rem;
        font-weight: 950;
        color: var(--tf-cyan);
        letter-spacing: -.035em;
    }
    .tf-hero-grid {
        display: grid;
        grid-template-columns: minmax(0, 1fr) 130px;
        gap: 1.2rem;
        align-items: center;
    }
    .tf-hero-icon {
        width: 118px;
        height: 118px;
        object-fit: contain;
        border-radius: 28px;
        filter: drop-shadow(0 18px 28px rgba(18,100,244,.16));
        justify-self: end;
    }
    .tf-sidebar-mini {
        display: flex;
        align-items: center;
        gap: .75rem;
        margin-bottom: .8rem;
    }
    .tf-sidebar-mini img {
        width: 58px;
        height: 58px;
        object-fit: contain;
        filter: drop-shadow(0 10px 18px rgba(18,100,244,.14));
    }
    .tf-sidebar-mini .name {
        line-height: 1.02;
        font-weight: 950;
        font-size: 1.35rem;
        letter-spacing: -.03em;
    }
    .tf-sidebar-mini .name span:first-child { color: var(--tf-navy); }
    .tf-sidebar-mini .name span:last-child { color: var(--tf-cyan); }
    @media (max-width: 900px) {
        .tf-hero-grid { grid-template-columns: 1fr; }
        .tf-hero-icon { justify-self: start; width: 88px; height: 88px; }
        .tf-topbar { flex-direction: column; align-items: stretch; }
    }


    .tf-login-panel {
        padding: 2rem 0 1rem 0;
    }
    .tf-login-title {
        font-size: 2.7rem;
        line-height: 1.12;
        margin: 1.3rem 0 .85rem 0;
        color: #10264b;
        font-weight: 950;
        letter-spacing: -.04em;
    }
    .tf-login-subtitle {
        color: #667799;
        font-size: 1.08rem;
        line-height: 1.65;
        max-width: 760px;
    }
    .tf-login-form-title {
        font-size: 1.75rem;
        color: #10264b;
        font-weight: 950;
        margin-bottom: 1rem;
    }
    .tf-clean-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 1rem;
        background: rgba(255,255,255,.78);
        border: 1px solid var(--tf-border);
        border-radius: 22px;
        padding: .9rem 1rem;
        margin: 0 0 1.1rem 0;
        box-shadow: 0 16px 38px rgba(16, 38, 75, .06);
    }
    .tf-clean-header-title {
        color: var(--tf-navy);
        font-weight: 950;
        letter-spacing: -.03em;
        font-size: 1.15rem;
    }
    .tf-clean-header-user {
        display: flex;
        align-items: center;
        justify-content: flex-end;
        gap: .7rem;
        color: var(--tf-navy);
        font-weight: 800;
        white-space: nowrap;
    }

    .tf-side-callout {
        margin: .95rem 0 1rem 0;
        padding: 1rem;
        border-radius: 18px;
        background: linear-gradient(135deg, rgba(0,184,200,.10), rgba(18,100,244,.07));
        border: 1px solid var(--tf-border);
        color: var(--tf-navy);
        box-shadow: 0 12px 26px rgba(16,38,75,.05);
    }
    .tf-mini-card {
        min-height: 154px;
        border-radius: 22px;
        padding: 1.05rem;
        background: rgba(255,255,255,.94);
        border: 1px solid var(--tf-border);
        box-shadow: 0 16px 38px rgba(16,38,75,.06);
        margin-bottom: .55rem;
    }
    .tf-mini-card p { color: var(--tf-muted); margin: .35rem 0 0 0; }
    .tf-mini-icon {
        width: 48px; height: 48px; border-radius: 16px;
        display:flex; align-items:center; justify-content:center;
        font-size:1.45rem; margin-bottom:.7rem;
        color:#fff; background: linear-gradient(135deg, #15c7bd, #00a9d6);
        box-shadow: 0 12px 24px rgba(16,38,75,.12);
    }
    .tf-mini-icon.purple { background: linear-gradient(135deg, #7c4dff, #b07cff); }
    .tf-mini-icon.orange { background: linear-gradient(135deg, #ff7a1a, #ffbd4a); }
    .tf-mini-icon.blue { background: linear-gradient(135deg, #246bfe, #2fa7ff); }

    </style>
    """, unsafe_allow_html=True)

def asset_data_uri(path: Path) -> str:
    """Retorna imagem local como data URI para uso seguro dentro de HTML/CSS."""
    try:
        if path.exists():
            data = base64.b64encode(path.read_bytes()).decode("ascii")
            return f"data:image/png;base64,{data}"
    except Exception:
        pass
    return ""


def crop_icon_from_full_logo_bytes() -> bytes:
    """Extrai somente o ícone do logo completo para evitar usar assets errados no app."""
    if Image is None or not LOGO_FULL.exists():
        return b""
    try:
        with Image.open(LOGO_FULL) as img:
            w, h = img.size
            # Recorta apenas a área do ícone à esquerda do logo principal.
            box = (max(0, int(w * 0.02)), max(0, int(h * 0.15)), max(1, int(w * 0.28)), max(1, int(h * 0.84)))
            cropped = img.crop(box)
            buf = io.BytesIO()
            cropped.save(buf, format="PNG")
            return buf.getvalue()
    except Exception:
        return b""


def brand_icon_data_uri() -> str:
    icon_bytes = crop_icon_from_full_logo_bytes()
    if icon_bytes:
        try:
            return f"data:image/png;base64,{base64.b64encode(icon_bytes).decode('ascii')}"
        except Exception:
            pass
    icon_uri = asset_data_uri(LOGO_ICON)
    if icon_uri:
        return icon_uri
    return asset_data_uri(LOGO_FULL)


def page_icon_object():
    if Image is not None:
        try:
            icon_bytes = crop_icon_from_full_logo_bytes()
            if icon_bytes:
                return Image.open(io.BytesIO(icon_bytes))
            if LOGO_ICON.exists():
                return Image.open(LOGO_ICON)
        except Exception:
            pass
    return "🎙️"


st.set_page_config(
    page_title="Transcreve Fácil",
    page_icon=page_icon_object(),
    layout="wide",
)


def brand_inline_html(compact: bool = False) -> str:
    icon_uri = brand_icon_data_uri()
    if icon_uri:
        if compact:
            return (
                f'<div class="tf-sidebar-mini"><img src="{icon_uri}" alt="Transcreve Fácil">'
                '<div class="name"><span>Transcreve</span><br><span>Fácil</span></div></div>'
            )
        return (
            f'<div class="tf-brand-row"><img src="{icon_uri}" alt="Transcreve Fácil">'
            '<div class="tf-brand-title"><div class="one">Transcreve</div><div class="two">Fácil</div></div></div>'
        )
    return (
        '<div class="tf-brand-title"><div class="one">Transcreve</div><div class="two">Fácil</div></div>'
    )


def show_icon(width=120):
    try:
        icon_uri = brand_icon_data_uri()
        if icon_uri:
            st.markdown(f"<img src='{icon_uri}' style='width:{width}px; height:auto; border-radius:20px; filter:drop-shadow(0 12px 24px rgba(18,100,244,.18));'>", unsafe_allow_html=True)
            return
    except Exception:
        pass
    st.markdown("<div style='font-size:3rem;'>🎙️</div>", unsafe_allow_html=True)


def show_logo(width=260):
    # Mantemos apenas o ícone + texto no layout para evitar que imagens antigas apareçam na tela inicial.
    st.markdown(brand_inline_html(), unsafe_allow_html=True)

AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac"}
VIDEO_EXTS = {".mp4", ".mov", ".avi", ".mkv", ".webm", ".mpeg", ".mpg"}
SUPPORTED_EXTS = sorted(AUDIO_EXTS | VIDEO_EXTS)
SUPPORTED_URL_DOMAINS = {"youtube.com", "www.youtube.com", "m.youtube.com", "youtu.be"}

# Limites conservadores para uso em Streamlit Cloud. O limite real do upload pode variar.
RECOMMENDED_MAX_MB = 150
RECOMMENDED_MAX_MINUTES = 45


# -------------------------
# Autenticação
# -------------------------
def get_secret_dict(section_name: str) -> dict:
    try:
        section = st.secrets.get(section_name, {})
        return dict(section)
    except Exception:
        return {}


def valid_institutional_email(email: str) -> bool:
    email = (email or "").strip().lower()
    if not email.endswith(ALLOWED_DOMAIN):
        return False
    return re.match(r"^[a-z0-9._%+\-]+@tre-ba\.jus\.br$", email) is not None


def authenticate(email: str, password: str) -> tuple[bool, str]:
    email = (email or "").strip().lower()
    password = password or ""

    if not valid_institutional_email(email):
        return False, "Use um e-mail institucional @tre-ba.jus.br."

    users = get_secret_dict("users")
    profiles = get_secret_dict("profiles")

    if users:
        expected_password = users.get(email)
        if expected_password and password == str(expected_password):
            return True, str(profiles.get(email, "usuario"))
        return False, "E-mail ou senha inválidos."

    # Primeiro acesso/fallback. Troque por Secrets assim que o app estiver publicado.
    if email == DEFAULT_USER and password == DEFAULT_PASSWORD:
        return True, "admin"

    try:
        app_password = st.secrets.get("APP_PASSWORD", DEFAULT_PASSWORD)
    except Exception:
        app_password = DEFAULT_PASSWORD
    if email == DEFAULT_USER and password == app_password:
        return True, "admin"

    return False, "Usuário não cadastrado. Configure os usuários em Secrets."


def login_screen():
    inject_css()
    st.markdown("<div class='tf-login-panel'>", unsafe_allow_html=True)
    left, right = st.columns([1.15, 0.85], gap="large")

    with left:
        st.markdown(brand_inline_html(), unsafe_allow_html=True)
        st.markdown(
            """
            <div class="tf-login-title">Sua central privada de transcrição</div>
            <div class="tf-login-subtitle">
                Transcreva vídeos e áudios, gere arquivos editáveis, fragmente mídias,
                compacte documentos e transforme conteúdo em material útil.
            </div>
            """,
            unsafe_allow_html=True,
        )

    with right:
        st.markdown("<div class='tf-login-form-title'>Acesso institucional</div>", unsafe_allow_html=True)
        with st.form("login_form"):
            email = st.text_input("E-mail institucional", placeholder="vmsoares@tre-ba.jus.br")
            password = st.text_input("Senha", type="password")
            submitted = st.form_submit_button("Entrar", use_container_width=True)

        if submitted:
            ok, result = authenticate(email, password)
            if ok:
                st.session_state["authenticated"] = True
                st.session_state["user_email"] = email.strip().lower()
                st.session_state["user_profile"] = result
                st.rerun()
            else:
                st.error(result)

        if not get_secret_dict("users"):
            with st.expander("Primeiro acesso"):
                st.code(f"E-mail: {DEFAULT_USER}\nSenha: {DEFAULT_PASSWORD}")
                st.caption("Depois, substitua por usuários configurados em Secrets.")

    st.markdown("</div>", unsafe_allow_html=True)


def logout_button():
    col1, col2 = st.columns([5, 1])
    with col1:
        st.caption(
            f"Usuário: {st.session_state.get('user_email', '')} | "
            f"Perfil: {st.session_state.get('user_profile', 'usuario')}"
        )
    with col2:
        if st.button("Sair", use_container_width=True):
            st.session_state.clear()
            st.rerun()


# -------------------------
# Utilidades
# -------------------------
def seconds_to_hhmmss(seconds: float) -> str:
    seconds = max(0, int(seconds or 0))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def safe_filename(name: str) -> str:
    stem = Path(name).stem[:80] or "transcricao"
    stem = re.sub(r"[^a-zA-Z0-9_.\-]+", "_", stem).strip("_")
    return stem or "transcricao"


def run_command(cmd: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, capture_output=True, text=True)


def get_duration_seconds(path: str) -> float | None:
    cmd = [
        "ffprobe", "-v", "error", "-show_entries", "format=duration",
        "-of", "default=noprint_wrappers=1:nokey=1", path,
    ]
    result = run_command(cmd)
    if result.returncode != 0:
        return None
    try:
        return float(result.stdout.strip())
    except Exception:
        return None


def extract_audio(input_path: str, output_path: str):
    cmd = [
        "ffmpeg", "-y", "-i", input_path,
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        output_path,
    ]
    result = run_command(cmd)
    if result.returncode != 0:
        raise RuntimeError(result.stderr[-2500:] if result.stderr else "Falha ao executar FFmpeg.")


def make_zip_from_paths(paths: list[str], zip_path: str):
    """Cria um ZIP com os arquivos informados sem recomprimir, para reduzir uso de CPU/memória."""
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_STORED, allowZip64=True) as zf:
        for file_path in paths:
            if os.path.exists(file_path):
                zf.write(file_path, arcname=os.path.basename(file_path))


def fragment_media_by_duration(input_path: str, output_dir: str, segment_seconds: int, original_suffix: str) -> list[str]:
    """Divide áudio/vídeo em partes por duração, preservando o formato quando possível."""
    suffix = original_suffix.lower() if original_suffix else ".mp4"
    pattern = os.path.join(output_dir, f"parte_%03d{suffix}")
    cmd = [
        "ffmpeg", "-y", "-i", input_path,
        "-map", "0", "-c", "copy",
        "-f", "segment", "-segment_time", str(int(segment_seconds)),
        "-reset_timestamps", "1", pattern,
    ]
    result = run_command(cmd)
    if result.returncode != 0:
        # Segunda tentativa, reencodando para MP4 quando o corte sem reencodar falhar.
        pattern = os.path.join(output_dir, "parte_%03d.mp4")
        cmd = [
            "ffmpeg", "-y", "-i", input_path,
            "-c:v", "libx264", "-preset", "veryfast", "-crf", "28",
            "-c:a", "aac", "-b:a", "96k",
            "-f", "segment", "-segment_time", str(int(segment_seconds)),
            "-reset_timestamps", "1", pattern,
        ]
        result = run_command(cmd)
        if result.returncode != 0:
            raise RuntimeError(result.stderr[-2500:] if result.stderr else "Falha ao fragmentar o arquivo com FFmpeg.")
    return sorted(str(x) for x in Path(output_dir).glob("parte_*.*"))


def split_file_by_size(input_path: str, output_dir: str, chunk_mb: int, original_name: str) -> list[str]:
    """Divide qualquer arquivo em partes binárias. Para usar depois, as partes precisam ser reunidas em ordem."""
    chunk_size = int(chunk_mb * 1024 * 1024)
    base = safe_filename(original_name)
    parts = []
    with open(input_path, "rb") as src:
        index = 1
        while True:
            data = src.read(chunk_size)
            if not data:
                break
            part_path = os.path.join(output_dir, f"{base}.part{index:03d}")
            with open(part_path, "wb") as dst:
                dst.write(data)
            parts.append(part_path)
            index += 1
    return parts


def compress_media(input_path: str, output_path: str, suffix: str, preset: str) -> str:
    """Compacta mídia por reencodação. Vídeo sai em MP4; áudio sai em MP3."""
    suffix = suffix.lower()
    if preset == "Alta qualidade":
        crf, audio_bitrate = "26", "128k"
    elif preset == "Equilibrado":
        crf, audio_bitrate = "30", "96k"
    else:
        crf, audio_bitrate = "34", "64k"

    if suffix in VIDEO_EXTS:
        cmd = [
            "ffmpeg", "-y", "-i", input_path,
            "-c:v", "libx264", "-preset", "veryfast", "-crf", crf,
            "-c:a", "aac", "-b:a", audio_bitrate,
            "-movflags", "+faststart", output_path,
        ]
    elif suffix in AUDIO_EXTS:
        cmd = [
            "ffmpeg", "-y", "-i", input_path,
            "-vn", "-codec:a", "libmp3lame", "-b:a", audio_bitrate,
            output_path,
        ]
    else:
        raise ValueError("Formato de mídia não suportado para compactação.")

    result = run_command(cmd)
    if result.returncode != 0:
        raise RuntimeError(result.stderr[-2500:] if result.stderr else "Falha ao compactar mídia com FFmpeg.")
    return output_path


def bytes_from_file(path: str) -> bytes:
    with open(path, "rb") as f:
        return f.read()


def is_supported_url(url: str) -> bool:
    try:
        parsed = urlparse((url or "").strip())
        host = (parsed.netloc or "").lower()
        return parsed.scheme in {"http", "https"} and host in SUPPORTED_URL_DOMAINS
    except Exception:
        return False


def friendly_youtube_error_message(error: Exception) -> str:
    """Converte erros técnicos do yt-dlp em orientação prática para o usuário."""
    msg = str(error)
    msg_lower = msg.lower()

    if "sign in to confirm" in msg_lower or "not a bot" in msg_lower or "cookies" in msg_lower:
        return (
            "O YouTube bloqueou o download automático neste servidor e pediu confirmação de que não é robô. "
            "Tente novamente usando o modo compatível. Se persistir, baixe o arquivo no seu computador e envie pelo upload. "
            "Como alternativa avançada, é possível anexar um arquivo cookies.txt exportado do navegador, mas use apenas em app privado."
        )
    if "private video" in msg_lower or "this video is private" in msg_lower:
        return "O vídeo parece ser privado. Use um vídeo público/autorizado ou envie o arquivo manualmente."
    if "video unavailable" in msg_lower or "unavailable" in msg_lower:
        return "O vídeo está indisponível para download automático. Confira o link ou envie o arquivo manualmente."
    if "copyright" in msg_lower:
        return "O vídeo possui restrição de direitos autorais ou disponibilidade. Use apenas conteúdo autorizado e, se necessário, envie o arquivo manualmente."
    if "unsupported url" in msg_lower:
        return "URL não suportada. Use links do YouTube nos formatos youtube.com/watch?v=... ou youtu.be/..."
    return (
        "Não foi possível baixar o áudio automaticamente. O YouTube pode ter bloqueado o servidor ou o vídeo pode ter restrições. "
        "Tente novamente ou envie o arquivo pelo upload."
    )


def write_uploaded_cookies(cookies_file, output_dir: str) -> str | None:
    """Salva um cookies.txt temporário, quando o usuário optar pelo modo avançado."""
    if cookies_file is None:
        return None
    cookies_path = os.path.join(output_dir, "cookies.txt")
    data = cookies_file.getvalue()
    if not data:
        return None
    with open(cookies_path, "wb") as f:
        f.write(data)
    return cookies_path


def youtube_base_options(base: str, strategy: str, cookies_path: str | None = None) -> dict:
    """Opções do yt-dlp com estratégias progressivas de compatibilidade."""
    common = {
        "format": "bestaudio[ext=m4a]/bestaudio/best",
        "outtmpl": base + ".%(ext)s",
        "noplaylist": True,
        "quiet": True,
        "no_warnings": True,
        "socket_timeout": 30,
        "retries": 3,
        "fragment_retries": 3,
        "extractor_retries": 3,
        "forceipv4": True,
        "geo_bypass": True,
        "nocheckcertificate": True,
        "cachedir": False,
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/126.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7",
        },
        "postprocessors": [
            {
                "key": "FFmpegExtractAudio",
                "preferredcodec": "wav",
            }
        ],
        "postprocessor_args": ["-ar", "16000", "-ac", "1"],
    }
    if cookies_path:
        common["cookiefile"] = cookies_path

    # Estratégias conhecidas que às vezes reduzem bloqueios do YouTube em nuvem.
    if strategy == "web":
        common["extractor_args"] = {"youtube": {"player_client": ["web"]}}
    elif strategy == "android":
        common["extractor_args"] = {"youtube": {"player_client": ["android"]}}
    elif strategy == "ios":
        common["extractor_args"] = {"youtube": {"player_client": ["ios"]}}
    elif strategy == "tv":
        common["extractor_args"] = {"youtube": {"player_client": ["tv"]}}
    elif strategy == "default":
        pass

    return common


def download_audio_from_url(url: str, output_dir: str, cookies_path: str | None = None, mode: str = "auto") -> tuple[str, dict]:
    """Baixa apenas o áudio de uma URL compatível e converte para WAV mono 16kHz.

    No Streamlit Cloud, o YouTube pode bloquear datacenters. Por isso a função tenta
    estratégias diferentes antes de desistir e retorna uma mensagem amigável.
    """
    url = (url or "").strip()
    if not is_supported_url(url):
        raise ValueError("URL não suportada. Use uma URL do YouTube, como youtube.com ou youtu.be.")

    # Ordem conservadora. Se houver cookies, tenta primeiro com modo padrão + cookies.
    if mode == "rapido":
        strategies = ["default", "web"]
    elif mode == "compatibilidade":
        strategies = ["web", "android", "ios", "tv", "default"]
    else:
        strategies = ["default", "web", "android", "ios", "tv"]

    last_error = None
    for idx, strategy in enumerate(strategies, start=1):
        base = os.path.join(output_dir, f"audio_youtube_{idx}_{strategy}")
        ydl_opts = youtube_base_options(base, strategy, cookies_path=cookies_path)
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)

            wav_path = base + ".wav"
            if not os.path.exists(wav_path):
                matches = list(Path(output_dir).glob(f"audio_youtube_{idx}_{strategy}*.wav"))
                if matches:
                    wav_path = str(matches[0])
                else:
                    raise RuntimeError("O áudio foi baixado, mas o arquivo WAV convertido não foi encontrado.")

            metadata = {
                "title": info.get("title") or "video_youtube",
                "duration": info.get("duration"),
                "webpage_url": info.get("webpage_url") or url,
                "uploader": info.get("uploader") or "",
                "download_strategy": strategy,
                "used_cookies": bool(cookies_path),
            }
            return wav_path, metadata
        except Exception as exc:
            last_error = exc
            continue

    message = friendly_youtube_error_message(last_error or RuntimeError("Falha desconhecida."))
    raise RuntimeError(message) from last_error


RUNTIME_PKG_DIR = Path(tempfile.gettempdir()) / "transcreve_facil_runtime_packages"


def ensure_runtime_package_path():
    """Permite instalar dependências em /tmp, sem escrever no venv do Streamlit Cloud."""
    RUNTIME_PKG_DIR.mkdir(parents=True, exist_ok=True)
    runtime_path = str(RUNTIME_PKG_DIR)
    if runtime_path not in sys.path:
        sys.path.insert(0, runtime_path)


def install_python_package(package_spec: str):
    """Instala pacote em pasta temporária, evitando PermissionError no Streamlit Cloud."""
    ensure_runtime_package_path()
    result = subprocess.run(
        [
            sys.executable,
            "-m",
            "pip",
            "install",
            "--no-cache-dir",
            "--disable-pip-version-check",
            "--target",
            str(RUNTIME_PKG_DIR),
            package_spec,
        ],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        details = (result.stderr or result.stdout or "").strip()
        raise RuntimeError(details[-4000:] if details else f"Falha ao instalar {package_spec}.")
    ensure_runtime_package_path()
    return True


@st.cache_resource(show_spinner=False)
def load_model(model_size: str):
    ensure_runtime_package_path()
    try:
        from faster_whisper import WhisperModel
    except Exception:
        # Instalação sob demanda. Pinamos versões para reduzir incompatibilidades no Streamlit Cloud.
        try:
            install_python_package("ctranslate2==4.6.0")
            install_python_package("faster-whisper==1.1.1")
        except Exception:
            # Segunda tentativa sem versões fixas, caso o ambiente tenha outra versão de Python.
            install_python_package("faster-whisper")
        from faster_whisper import WhisperModel
    return WhisperModel(model_size, device="cpu", compute_type="int8")


def transcribe_with_google_fallback(audio_path: str, duration: float | None, include_timestamps: bool, progress):
    """Fallback simples quando faster-whisper não instala no Streamlit Cloud.

    Observação: este fallback usa o serviço gratuito do Google Web Speech via biblioteca
    SpeechRecognition. Portanto, use apenas se aceitar esse processamento externo.
    """
    ensure_runtime_package_path()
    try:
        import speech_recognition as sr
    except Exception:
        install_python_package("SpeechRecognition")
        import speech_recognition as sr

    wav_path = audio_path
    if not audio_path.lower().endswith(".wav"):
        wav_path = audio_path + ".fallback.wav"
        extract_audio(audio_path, wav_path)

    recognizer = sr.Recognizer()
    chunk_seconds = 45
    lines = []
    plain_parts = []
    segments_data = []

    with sr.AudioFile(wav_path) as source:
        total = float(getattr(source, "DURATION", None) or duration or 0)
        offset = 0.0
        while True:
            audio = recognizer.record(source, duration=chunk_seconds)
            if not getattr(audio, "frame_data", b""):
                break

            start = offset
            end = offset + chunk_seconds
            offset = end

            try:
                text = recognizer.recognize_google(audio, language="pt-BR")
            except sr.UnknownValueError:
                text = ""
            except Exception as exc:
                text = f"[trecho não reconhecido: {exc}]"

            text = (text or "").strip()
            if text:
                plain_parts.append(text)
                segments_data.append({"start": start, "end": end, "text": text})
                if include_timestamps:
                    lines.append(f"[{seconds_to_hhmmss(start)} - {seconds_to_hhmmss(end)}] {text}")
                else:
                    lines.append(text)

            if total:
                pct = 55 + min(40, int((min(offset, total) / total) * 40))
                progress.progress(pct, text=f"Transcrevendo com fallback... {min(100, int((min(offset, total) / total) * 100))}%")

            if total and offset >= total:
                break

    if not lines:
        raise RuntimeError("Nenhuma fala foi reconhecida pelo fallback.")

    return "\\n".join(lines), "\\n".join(plain_parts), segments_data, "google-web-speech-fallback", None


def transcribe_audio_engine(audio_path: str, duration: float | None, model_size: str, beam_size: int, include_timestamps: bool, progress):
    """Tenta faster-whisper; se falhar, usa fallback online simples."""
    try:
        progress.progress(40, text="Carregando modelo de transcrição...")
        model = load_model(model_size)

        progress.progress(55, text="Transcrevendo. Aguarde...")
        segments, info = model.transcribe(
            audio_path,
            language="pt",
            vad_filter=True,
            beam_size=beam_size,
        )

        lines = []
        plain_parts = []
        segments_data = []
        last_progress = 55

        for seg in segments:
            text = (seg.text or "").strip()
            if not text:
                continue
            segments_data.append({"start": seg.start, "end": seg.end, "text": text})
            plain_parts.append(text)
            if include_timestamps:
                lines.append(f"[{seconds_to_hhmmss(seg.start)} - {seconds_to_hhmmss(seg.end)}] {text}")
            else:
                lines.append(text)

            if duration and duration > 0:
                pct = 55 + min(40, int((seg.end / duration) * 40))
                if pct > last_progress:
                    progress.progress(pct, text=f"Transcrevendo... {min(100, int((seg.end / duration) * 100))}%")
                    last_progress = pct

        if not lines:
            raise RuntimeError("Nenhuma fala foi identificada no arquivo.")

        return "\\n".join(lines), "\\n".join(plain_parts), segments_data, "faster-whisper", info

    except Exception as fast_error:
        st.warning(
            "O motor local faster-whisper não pôde ser instalado/carregado no ambiente principal. "
            "Vou tentar um fallback online simples para concluir a transcrição."
        )
        with st.expander("Detalhes técnicos do motor local"):
            st.code(str(fast_error)[-4000:])
        return transcribe_with_google_fallback(audio_path, duration, include_timestamps, progress)



def save_docx(lines: list[str], metadata: dict) -> bytes:
    doc = Document()
    doc.add_heading("Transcrição", level=1)
    doc.add_paragraph(f"Sistema: {APP_NAME}")
    doc.add_paragraph(f"Gerada em: {metadata.get('generated_at', '')}")
    if metadata.get("filename"):
        doc.add_paragraph(f"Arquivo: {metadata['filename']}")
    if metadata.get("duration"):
        doc.add_paragraph(f"Duração estimada: {metadata['duration']}")
    doc.add_paragraph("")

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for line in lines:
        doc.add_paragraph(line)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    with open(tmp.name, "rb") as f:
        data = f.read()
    os.unlink(tmp.name)
    return data


def save_pdf(lines: list[str], metadata: dict) -> bytes:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp.close()
    doc = SimpleDocTemplate(
        tmp.name,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "NormalWrapped",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    story = []
    story.append(Paragraph("Transcrição", styles["Title"]))
    story.append(Paragraph(f"Sistema: {APP_NAME}", normal))
    story.append(Paragraph(f"Gerada em: {metadata.get('generated_at', '')}", normal))
    if metadata.get("filename"):
        story.append(Paragraph(f"Arquivo: {metadata['filename']}", normal))
    if metadata.get("duration"):
        story.append(Paragraph(f"Duração estimada: {metadata['duration']}", normal))
    story.append(Spacer(1, 12))

    for line in lines:
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe_line, normal))

    doc.build(story)
    with open(tmp.name, "rb") as f:
        data = f.read()
    os.unlink(tmp.name)
    return data


def build_srt(segments_data: list[dict]) -> str:
    def srt_time(seconds: float) -> str:
        seconds = max(0, float(seconds or 0))
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        ms = int((seconds - int(seconds)) * 1000)
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    blocks = []
    for i, seg in enumerate(segments_data, start=1):
        blocks.append(
            f"{i}\n{srt_time(seg['start'])} --> {srt_time(seg['end'])}\n{seg['text']}"
        )
    return "\n\n".join(blocks)


def build_prompts(plain_text: str) -> dict:
    return {
        "Revisar transcrição": (
            "Revise a transcrição abaixo, corrigindo pontuação, quebras de parágrafo "
            "e termos evidentemente incorretos, sem alterar o sentido. Preserve datas, números, "
            "nomes próprios e fundamentos citados:\n\n" + plain_text
        ),
        "Resumo objetivo": (
            "Faça um resumo objetivo da transcrição abaixo, destacando os principais pontos, "
            "decisões, orientações, datas relevantes e eventuais ressalvas:\n\n" + plain_text
        ),
        "Ata de reunião": (
            "Transforme a transcrição abaixo em uma ata formal, com pauta, participantes quando "
            "identificáveis, deliberações, encaminhamentos, pendências e observações finais:\n\n" + plain_text
        ),
        "Tabela prática": (
            "Transforme a transcrição abaixo em uma tabela prática, com colunas de tema, orientação, "
            "fundamento citado, providência e observações:\n\n" + plain_text
        ),
        "Checklist": (
            "Transforme a transcrição abaixo em um checklist de providências, separando o que deve ser feito, "
            "responsável sugerido, prazo se houver e observações:\n\n" + plain_text
        ),
        "Material de estudo": (
            "Transforme a transcrição abaixo em material de estudo, com tópicos, explicações, conceitos-chave, "
            "quadros comparativos e pontos de atenção para prova ou aplicação prática:\n\n" + plain_text
        ),
    }


def estimate_warning(file_mb: float, duration: float | None, model_size: str):
    if file_mb > RECOMMENDED_MAX_MB:
        st.warning(
            f"Arquivo com {file_mb:.1f} MB. No Streamlit Cloud, arquivos acima de "
            f"{RECOMMENDED_MAX_MB} MB podem demorar ou falhar."
        )
    if duration and duration / 60 > RECOMMENDED_MAX_MINUTES:
        st.warning(
            f"Duração estimada de {duration / 60:.1f} min. Para arquivos longos, a versão local tende a ser melhor."
        )
    if model_size != "small":
        st.warning("No Streamlit Cloud, o modelo small é o mais estável. Medium/large podem ficar lentos ou travar.")



def windows_safe_bat_text(text: str) -> str:
    text = (text or "").strip()
    return text.replace('"', '').replace('\n', '').replace('\r', '')


def build_youtube_helper_bat(url: str = "", mode: str = "audio_mp3") -> str:
    """Gera um arquivo .bat assistido para baixar conteudo autorizado do YouTube no Windows."""
    url = windows_safe_bat_text(url)
    ask_url = not bool(url)

    if mode == "video_mp4":
        title = "Transcreve Facil - YouTube para video"
        main_cmd = '%PY% -m yt_dlp -f "bv*+ba/best" --merge-output-format mp4 --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        fallback_cmd = '%PY% -m yt_dlp -f "best" --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        success_msg = "Video baixado."
    elif mode == "audio_original":
        title = "Transcreve Facil - YouTube para audio original"
        main_cmd = '%PY% -m yt_dlp -f "bestaudio/best" --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        fallback_cmd = '%PY% -m yt_dlp -f "best" --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        success_msg = "Audio baixado."
    else:
        title = "Transcreve Facil - YouTube para MP3"
        main_cmd = '%PY% -m yt_dlp -x --audio-format mp3 --audio-quality 0 --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        fallback_cmd = '%PY% -m yt_dlp -f "bestaudio/best" --paths "%OUTDIR%" -o "TranscreveFacil_%(title).80s.%(ext)s" "%URL%"'
        success_msg = "Audio preparado. Se o MP3 falhar por falta de FFmpeg, o script tenta baixar o audio original."

    set_url_line = 'set /p URL=Cole a URL do YouTube e pressione ENTER: ' if ask_url else f'set "URL={url}"'

    bat = f"""@echo off
chcp 65001 >nul
setlocal
title {title}

echo ============================================================
echo  Transcreve Facil - Conversao Privada
echo ============================================================
echo.
echo Use apenas em videos seus, autorizados ou permitidos.
echo O arquivo sera salvo em: Downloads\\TranscreveFacil
echo.

{set_url_line}

if "%URL%"=="" (
    echo Nenhuma URL informada.
    pause
    exit /b 1
)

set "OUTDIR=%USERPROFILE%\\Downloads\\TranscreveFacil"
if not exist "%OUTDIR%" mkdir "%OUTDIR%"

echo.
echo Verificando Python...
where py >nul 2>nul
if %errorlevel%==0 (
    set "PY=py"
) else (
    set "PY=python"
)

%PY% --version >nul 2>nul
if errorlevel 1 (
    echo Python nao encontrado.
    echo Instale o Python em https://www.python.org/downloads/ e marque "Add Python to PATH".
    pause
    exit /b 1
)

echo.
echo Atualizando ferramenta de download...
%PY% -m pip install -U yt-dlp

echo.
echo Baixando conteudo autorizado...
{main_cmd}

if errorlevel 1 (
    echo.
    echo A primeira tentativa falhou. Tentando modo alternativo...
    {fallback_cmd}
)

if errorlevel 1 (
    echo.
    echo Nao foi possivel baixar o conteudo.
    echo O YouTube pode ter bloqueado o acesso, exigido login ou o video pode estar indisponivel.
    pause
    exit /b 1
)

echo.
echo {success_msg}
echo Abra a pasta, selecione o arquivo gerado e envie no Transcreve Facil pela aba Transcrever.
echo.
start "" "%OUTDIR%"
pause
"""
    return bat.replace("\n", "\r\n")


def render_conversao_privada_page():
    st.subheader("Conversão privada / YouTube assistido")
    st.info(
        "Este módulo resolve o bloqueio do YouTube no Streamlit Cloud. "
        "Você gera um pequeno baixador local pelo próprio sistema, executa no Windows, "
        "e depois envia o arquivo gerado pela aba Transcrever."
    )

    st.markdown("### 1. Cole a URL do vídeo")
    url = st.text_input(
        "URL do YouTube",
        placeholder="https://www.youtube.com/watch?v=...",
        key="conv_priv_url",
    )

    col_a, col_b, col_c = st.columns(3)
    with col_a:
        st.markdown(
            "<div class='tf-mini-card'><div class='tf-mini-icon'>🎧</div><b>YouTube para MP3</b>"
            "<p>Melhor opção para transcrição. Tenta converter para MP3.</p></div>",
            unsafe_allow_html=True,
        )
        st.download_button(
            "Baixar assistente MP3",
            data=build_youtube_helper_bat(url, "audio_mp3"),
            file_name="TranscreveFacil_YouTube_para_MP3.bat",
            mime="application/octet-stream",
            use_container_width=True,
        )
    with col_b:
        st.markdown(
            "<div class='tf-mini-card'><div class='tf-mini-icon blue'>🎙️</div><b>Áudio original</b>"
            "<p>Baixa o melhor áudio disponível, sem exigir conversão.</p></div>",
            unsafe_allow_html=True,
        )
        st.download_button(
            "Baixar assistente áudio",
            data=build_youtube_helper_bat(url, "audio_original"),
            file_name="TranscreveFacil_YouTube_audio_original.bat",
            mime="application/octet-stream",
            use_container_width=True,
        )
    with col_c:
        st.markdown(
            "<div class='tf-mini-card'><div class='tf-mini-icon orange'>🎬</div><b>Vídeo MP4</b>"
            "<p>Baixa o vídeo para guardar ou processar depois.</p></div>",
            unsafe_allow_html=True,
        )
        st.download_button(
            "Baixar assistente vídeo",
            data=build_youtube_helper_bat(url, "video_mp4"),
            file_name="TranscreveFacil_YouTube_video_MP4.bat",
            mime="application/octet-stream",
            use_container_width=True,
        )

    st.markdown("### 2. Como usar")
    st.write(
        "1. Baixe uma das opções acima.\n\n"
        "2. Dê dois cliques no arquivo `.bat` baixado.\n\n"
        "3. Se o Windows mostrar aviso, escolha **Mais informações** e depois **Executar assim mesmo**, se confiar no arquivo.\n\n"
        "4. O arquivo será salvo em `Downloads > TranscreveFacil`.\n\n"
        "5. Volte ao sistema e envie o áudio pela aba **Transcrever**."
    )

    st.warning(
        "Use apenas em vídeos seus, autorizados ou permitidos. "
        "O arquivo .bat roda no seu computador porque o navegador não pode executar esse tipo de tarefa diretamente por segurança."
    )

    st.markdown("### 3. Assistente genérico")
    st.caption("Use este caso queira baixar o assistente uma vez e colar URLs diferentes quando ele abrir.")
    st.download_button(
        "Baixar assistente genérico de MP3",
        data=build_youtube_helper_bat("", "audio_mp3"),
        file_name="TranscreveFacil_Assistente_Generico_MP3.bat",
        mime="application/octet-stream",
        use_container_width=True,
    )




def build_local_fragmenter_bat() -> str:
    """Gera assistente local robusto usando PowerShell para evitar problemas de aspas/caminho no Windows."""
    ps_script = """
$ErrorActionPreference = "Stop"
Write-Host "============================================================"
Write-Host " Transcreve Facil - Fragmentador Local Robusto"
Write-Host "============================================================"
Write-Host ""
Write-Host "Use este assistente para arquivos grandes demais para o Streamlit Cloud."
Write-Host ""

$inputPath = Read-Host "Cole o caminho completo do arquivo ou arraste o arquivo aqui"
$inputPath = $inputPath.Trim('"')

if ([string]::IsNullOrWhiteSpace($inputPath) -or -not (Test-Path $inputPath)) {
    Write-Host "Arquivo nao encontrado."
    Read-Host "Pressione ENTER para sair"
    exit 1
}

$minutesText = Read-Host "Duracao de cada parte em minutos [padrao 10]"
if ([string]::IsNullOrWhiteSpace($minutesText)) { $minutes = 10 } else { $minutes = [int]$minutesText }
$seconds = $minutes * 60

$downloads = Join-Path $env:USERPROFILE "Downloads"
$outDir = Join-Path $downloads "TranscreveFacil_Fragmentado"
$ffRoot = Join-Path $downloads "TranscreveFacil_FFmpeg"
$ffZip = Join-Path $downloads "TranscreveFacil_FFmpeg.zip"

New-Item -ItemType Directory -Force -Path $outDir | Out-Null
New-Item -ItemType Directory -Force -Path $ffRoot | Out-Null

Write-Host ""
Write-Host "Procurando FFmpeg..."
$ffmpegCmd = Get-Command ffmpeg -ErrorAction SilentlyContinue
if ($ffmpegCmd) { $ffmpeg = $ffmpegCmd.Source } else { $ffmpeg = $null }

if (-not $ffmpeg) {
    $existing = Get-ChildItem -Path $ffRoot -Filter ffmpeg.exe -Recurse -ErrorAction SilentlyContinue | Select-Object -First 1
    if ($existing) { $ffmpeg = $existing.FullName }
}

if (-not $ffmpeg) {
    Write-Host "FFmpeg nao encontrado. Baixando versao portatil..."
    $url = "https://www.gyan.dev/ffmpeg/builds/ffmpeg-release-essentials.zip"
    [Net.ServicePointManager]::SecurityProtocol = [Net.SecurityProtocolType]::Tls12
    Invoke-WebRequest -Uri $url -OutFile $ffZip -UseBasicParsing
    Write-Host "Extraindo FFmpeg..."
    Expand-Archive -Path $ffZip -DestinationPath $ffRoot -Force
    $existing = Get-ChildItem -Path $ffRoot -Filter ffmpeg.exe -Recurse -ErrorAction SilentlyContinue | Select-Object -First 1
    if ($existing) { $ffmpeg = $existing.FullName }
}

if (-not $ffmpeg -or -not (Test-Path $ffmpeg)) {
    Write-Host "Nao foi possivel localizar ffmpeg.exe."
    Read-Host "Pressione ENTER para sair"
    exit 1
}

Write-Host ""
Write-Host ("FFmpeg localizado: " + $ffmpeg)
Write-Host "Fragmentando sem recodificar..."
$patternMp4 = Join-Path $outDir "parte_%03d.mp4"
$args1 = @("-hide_banner", "-y", "-i", $inputPath, "-map", "0", "-c", "copy", "-f", "segment", "-segment_time", "$seconds", "-reset_timestamps", "1", $patternMp4)
& $ffmpeg @args1
$code = $LASTEXITCODE

if ($code -ne 0) {
    Write-Host ""
    Write-Host "A divisao sem recodificar falhou. Tentando gerar audios MP3 leves..."
    $patternMp3 = Join-Path $outDir "audio_parte_%03d.mp3"
    $args2 = @("-hide_banner", "-y", "-i", $inputPath, "-vn", "-acodec", "libmp3lame", "-ar", "16000", "-ac", "1", "-b:a", "64k", "-f", "segment", "-segment_time", "$seconds", "-reset_timestamps", "1", $patternMp3)
    & $ffmpeg @args2
    $code = $LASTEXITCODE
}

if ($code -ne 0) {
    Write-Host "Nao foi possivel fragmentar o arquivo."
    Write-Host "Tente mover o arquivo para Downloads, reduzir o nome do arquivo ou escolher partes menores."
    Read-Host "Pressione ENTER para sair"
    exit 1
}

Write-Host ""
Write-Host ("Concluido. As partes foram salvas em: " + $outDir)
Start-Process explorer.exe $outDir
Read-Host "Pressione ENTER para sair"
"""
    encoded = ps_script.encode("utf-16le")
    import base64 as _b64
    ps_b64 = _b64.b64encode(encoded).decode("ascii")
    bat = f"""@echo off
chcp 65001 >nul
title Transcreve Facil - Fragmentador Local Robusto
powershell -NoProfile -ExecutionPolicy Bypass -EncodedCommand {ps_b64}
pause
"""
    return bat.replace("\n", "\r\n")


def show_large_file_local_fragmenter_notice():
    st.error(
        "Arquivo grande demais para fragmentar com segurança dentro do Streamlit Cloud. "
        "Para evitar queda do app no final do processamento, use o Fragmentador Local abaixo."
    )
    st.download_button(
        "Baixar Fragmentador Local robusto do Transcreve Fácil",
        data=build_local_fragmenter_bat(),
        file_name="TranscreveFacil_Fragmentador_Local.bat",
        mime="application/octet-stream",
        use_container_width=True,
    )
    st.info(
        "Depois de fragmentar no computador, envie as partes menores pela aba Transcrever. "
        "Recomendação: partes de 10 a 15 minutos."
    )

# -------------------------
# Tela principal
# -------------------------
def app_screen():
    inject_css()

    if "tf_page" not in st.session_state:
        st.session_state["tf_page"] = "Inicio"

    def go_to(page_name: str):
        st.session_state["tf_page"] = page_name

    user_email = st.session_state.get("user_email", "")

    with st.sidebar:
        st.markdown(brand_inline_html(compact=True), unsafe_allow_html=True)
        st.markdown("### Navegação")

        nav_items = [
            ("Inicio", "🏠", "Início"),
            ("Transcrever", "🎙️", "Transcrever"),
            ("Resultado", "📥", "Resultado"),
            ("Prompts", "✨", "Prompts"),
            ("Ferramentas", "🧰", "Ferramentas"),
            ("Conversao privada", "🔐", "Conversão privada"),
            ("YouTube local", "▶️", "YouTube local"),
            ("Ajuda", "❔", "Ajuda"),
        ]
        for page_key, icon, label in nav_items:
            selected = st.session_state.get("tf_page") == page_key
            if st.button(f"{icon}  {label}", key=f"nav_{page_key}", type="primary" if selected else "secondary", use_container_width=True):
                go_to(page_key)
                st.rerun()

        st.markdown("<div class='tf-side-callout'>🚀 <b>Transcreva mais, organize melhor e economize tempo.</b></div>", unsafe_allow_html=True)
        st.divider()
        st.header("Configurações")
        model_size = st.selectbox("Modelo", ["small", "medium", "large-v3"], index=0)
        include_timestamps = st.checkbox("Incluir marcação de tempo no texto", value=True)
        beam_size = st.slider("Precisão da busca", min_value=1, max_value=5, value=1)
        st.caption("Use 1 para mais velocidade. Use 5 para tentar melhorar a qualidade, com mais demora.")

        st.divider()
        st.header("Limites recomendados")
        st.write(f"Arquivo: até {RECOMMENDED_MAX_MB} MB")
        st.write(f"Duração: até {RECOMMENDED_MAX_MINUTES} min")
        st.write("Modelo online: small")

        if st.button("🧹 Limpar resultado atual", use_container_width=True):
            for key in ["last_transcription", "last_plain_text", "last_segments", "last_metadata"]:
                st.session_state.pop(key, None)
            st.success("Resultado limpo.")

    st.markdown(
        f"""
        <div class="tf-clean-header">
            <div class="tf-clean-header-title">Transcreve Fácil</div>
            <div class="tf-clean-header-user">
                <span class="tf-avatar">{(user_email[:1] or 'V').upper()}</span>
                <span>{user_email}</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    page = st.session_state.get("tf_page", "Inicio")

    if page == "Inicio":
        hero_icon_uri = brand_icon_data_uri()
        hero_icon_html = f'<img class="tf-hero-icon" src="{hero_icon_uri}" alt="Transcreve Fácil">' if hero_icon_uri else '<div class="tf-hero-icon" style="font-size:4rem; display:flex; align-items:center; justify-content:center;">🎙️</div>'
        st.markdown(
            f"""
            <div class="tf-hero">
                <div class="tf-hero-grid">
                    <div>
                        <h1>Transcritor de Vídeos e Áudios</h1>
                        <p>Envie arquivos, fragmente mídias, compacte documentos e gere transcrições em TXT, Word, PDF e SRT.</p>
                    </div>
                    {hero_icon_html}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.caption(APP_VERSION)

        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("<div class='tf-card'><div class='tf-icon teal'>📄</div><h3>Arquivos processados</h3><p>Transcreva vídeos, áudios e gere downloads editáveis.</p></div>", unsafe_allow_html=True)
        with c2:
            st.markdown("<div class='tf-card'><div class='tf-icon blue'>⏱️</div><h3>Tempo economizado</h3><p>Transforme aulas, reuniões e treinamentos em texto pesquisável.</p></div>", unsafe_allow_html=True)
        with c3:
            st.markdown("<div class='tf-card'><div class='tf-icon orange'>🧰</div><h3>Ferramentas integradas</h3><p>Fragmentador, compactador, PDF, Word, TXT e legendas SRT.</p></div>", unsafe_allow_html=True)

        st.markdown("### Ferramentas rápidas")
        q1, q2, q3, q4 = st.columns(4)
        with q1:
            st.markdown("<div class='tf-mini-card'><div class='tf-mini-icon'>🎙️</div><b>Nova transcrição</b><p>Enviar vídeo ou áudio.</p></div>", unsafe_allow_html=True)
            if st.button("Abrir Transcrever", key="quick_transcrever", use_container_width=True):
                go_to("Transcrever"); st.rerun()
        with q2:
            st.markdown("<div class='tf-mini-card'><div class='tf-mini-icon purple'>✂️</div><b>Fragmentador</b><p>Dividir arquivos grandes.</p></div>", unsafe_allow_html=True)
            if st.button("Abrir Ferramentas", key="quick_ferramentas", use_container_width=True):
                go_to("Ferramentas"); st.rerun()
        with q3:
            st.markdown("<div class='tf-mini-card'><div class='tf-mini-icon orange'>📥</div><b>Resultado</b><p>Baixar TXT, Word, PDF e SRT.</p></div>", unsafe_allow_html=True)
            if st.button("Ver Resultado", key="quick_resultado", use_container_width=True):
                go_to("Resultado"); st.rerun()
        with q4:
            st.markdown("<div class='tf-mini-card'><div class='tf-mini-icon blue'>✨</div><b>Prompts</b><p>Resumo, ata e tabela prática.</p></div>", unsafe_allow_html=True)
            if st.button("Gerar Prompts", key="quick_prompts", use_container_width=True):
                go_to("Prompts"); st.rerun()

        st.markdown("<div class='tf-card'><h3>Fluxo recomendado</h3><p><b>1.</b> Transcrever → <b>2.</b> Resultado → <b>3.</b> Prompts ou Ferramentas. Para YouTube, use primeiro o modo local quando houver bloqueio.</p></div>", unsafe_allow_html=True)
        return

    if page == "Transcrever":
        st.warning(
            "Use URLs apenas para vídeos seus, autorizados ou com permissão de uso. "
            "No Streamlit Cloud, o YouTube pode bloquear downloads automáticos. O upload manual continua sendo o caminho mais estável."
        )
        st.markdown(
            "<div class='tf-card'><h3>Enviar arquivo para transcrição</h3>"
            "<p>Use o botão de upload abaixo. O quadro anterior foi removido porque era apenas visual.</p>"
            "<p><b>Formatos aceitos:</b> MP3, MP4, WAV, M4A, MOV, AAC e outros formatos compatíveis.</p></div>",
            unsafe_allow_html=True,
        )

        origem = st.radio(
            "Fonte do conteúdo",
            ["Enviar arquivo", "URL do YouTube"],
            horizontal=True,
        )

        uploaded = None
        youtube_url = ""
        yt_mode = "auto"
        cookies_file = None
        if origem == "Enviar arquivo":
            uploaded = st.file_uploader(
                "Clique aqui para selecionar ou arraste o arquivo para esta área",
                type=[ext.replace(".", "") for ext in SUPPORTED_EXTS],
                help="Formatos aceitos: áudio e vídeo. Para vídeos, o sistema extrai o áudio automaticamente.",
                key="main_upload_real",
            )
            if uploaded is None:
                st.info("Nenhum arquivo selecionado ainda. Clique no campo acima ou arraste o arquivo para ele.")
            else:
                st.success(f"Arquivo selecionado: {uploaded.name}")
            ready_to_transcribe = uploaded is not None
        else:
            youtube_url = st.text_input(
                "Cole a URL do YouTube",
                placeholder="https://www.youtube.com/watch?v=... ou https://youtu.be/...",
            )
            yt_mode = st.selectbox(
                "Modo de download",
                ["auto", "compatibilidade", "rapido"],
                index=0,
                help="Use compatibilidade quando o YouTube bloquear o modo automático. Pode demorar mais."
            )
            with st.expander("Opção avançada: cookies.txt do navegador"):
                st.write(
                    "Use somente se o YouTube bloquear o download e apenas em app privado. "
                    "O arquivo é usado temporariamente durante a transcrição e não é salvo no repositório."
                )
                cookies_file = st.file_uploader(
                    "Anexar cookies.txt (opcional)",
                    type=["txt"],
                    help="Exportado do navegador por extensão própria. Não compartilhe esse arquivo com terceiros."
                )
            st.caption("Recurso experimental no Streamlit Cloud. Se aparecer bloqueio 403, not-a-bot ou login obrigatório, use a aba YouTube local.")
            ready_to_transcribe = bool(youtube_url.strip())

        if not ready_to_transcribe:
            st.info("Envie um arquivo ou cole uma URL do YouTube para começar.")
        else:
            if origem == "Enviar arquivo":
                file_mb = uploaded.size / (1024 * 1024)
                suffix = Path(uploaded.name).suffix.lower()
                display_name = uploaded.name
                display_type = "Vídeo" if suffix in VIDEO_EXTS else "Áudio"
                st.success(f"Arquivo carregado: {uploaded.name} ({file_mb:.1f} MB)")
            else:
                file_mb = 0.0
                suffix = ".url"
                display_name = youtube_url.strip()
                display_type = "YouTube"
                st.success("URL informada. O áudio será baixado ao iniciar a transcrição.")

            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.metric("Tamanho", f"{file_mb:.1f} MB" if origem == "Enviar arquivo" else "URL")
            with col_b:
                st.metric("Tipo", display_type)
            with col_c:
                st.metric("Modelo", model_size)

            if origem == "Enviar arquivo" and suffix not in AUDIO_EXTS and suffix not in VIDEO_EXTS:
                st.error("Formato não suportado.")
                return
            if origem == "URL do YouTube" and not is_supported_url(youtube_url):
                st.error("URL não suportada. Use uma URL do YouTube, como youtube.com ou youtu.be.")
                return

            if st.button("Transcrever", type="primary", use_container_width=True):
                progress = st.progress(0, text="Preparando...")
                status_box = st.empty()

                with tempfile.TemporaryDirectory() as tmpdir:
                    source_title = display_name
                    yt_meta = {}
                    if origem == "Enviar arquivo":
                        input_path = os.path.join(tmpdir, safe_filename(uploaded.name) + suffix)
                        with open(input_path, "wb") as f:
                            f.write(uploaded.getbuffer())

                        progress.progress(10, text="Arquivo salvo temporariamente.")
                        duration = get_duration_seconds(input_path)
                        if duration:
                            status_box.info(f"Duração estimada: {seconds_to_hhmmss(duration)}")
                        estimate_warning(file_mb, duration, model_size)

                        audio_path = input_path
                        try:
                            if suffix in VIDEO_EXTS:
                                progress.progress(25, text="Extraindo áudio do vídeo...")
                                audio_path = os.path.join(tmpdir, "audio_extraido.wav")
                                extract_audio(input_path, audio_path)
                            else:
                                progress.progress(25, text="Arquivo de áudio identificado.")
                        except Exception:
                            st.error("Não foi possível extrair o áudio do arquivo enviado.")
                            st.info("Confira se o arquivo não está corrompido. No Streamlit Cloud, o FFmpeg deve estar listado no packages.txt.")
                            return
                    else:
                        try:
                            progress.progress(15, text="Baixando áudio da URL. Se o YouTube bloquear, o sistema tentará modos alternativos...")
                            cookies_path = write_uploaded_cookies(cookies_file, tmpdir)
                            audio_path, yt_meta = download_audio_from_url(youtube_url, tmpdir, cookies_path=cookies_path, mode=yt_mode)
                            source_title = yt_meta.get("title") or "video_youtube"
                            duration = yt_meta.get("duration") or get_duration_seconds(audio_path)
                            if duration:
                                status_box.info(f"Duração estimada: {seconds_to_hhmmss(duration)}")
                            estimate_warning(file_mb, duration, model_size)
                            progress.progress(30, text="Áudio baixado e convertido.")
                        except Exception as e:
                            st.error("Não foi possível baixar o áudio diretamente pelo Streamlit Cloud.")
                            st.warning(str(e))
                            st.info(
                                "Caminho mais estável: abra a aba YouTube local, gere o comando para baixar o áudio no seu computador "
                                "e depois envie o MP3 pela opção Enviar arquivo."
                            )
                            return

                    try:
                        final_text, plain_text, segments_data, engine_name, info = transcribe_audio_engine(
                            audio_path=audio_path,
                            duration=duration,
                            model_size=model_size,
                            beam_size=beam_size,
                            include_timestamps=include_timestamps,
                            progress=progress,
                        )

                        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
                        metadata = {
                            "filename": source_title,
                            "file_mb": f"{file_mb:.1f} MB" if origem == "Enviar arquivo" else "URL",
                            "duration": seconds_to_hhmmss(duration) if duration else "não identificada",
                            "model": model_size,
                            "source": origem,
                            "url": yt_meta.get("webpage_url", "") if origem == "URL do YouTube" else "",
                            "uploader": yt_meta.get("uploader", "") if origem == "URL do YouTube" else "",
                            "generated_at": generated_at,
                            "language_probability": getattr(info, "language_probability", None) if info is not None else None,
                            "engine": engine_name,
                        }

                        st.session_state["last_transcription"] = final_text
                        st.session_state["last_plain_text"] = plain_text
                        st.session_state["last_segments"] = segments_data
                        st.session_state["last_metadata"] = metadata

                        progress.progress(100, text="Transcrição concluída.")
                        status_box.success("Transcrição concluída. Abra a aba Resultado para baixar os arquivos.")
                    except Exception as e:
                        st.error("Erro durante a transcrição.")
                        st.info("Tente novamente com arquivo menor, envie apenas o áudio em MP3/WAV ou use a Conversão privada para preparar o arquivo.")
                        st.caption(str(e))
                        return

    elif page == "Resultado":
        if not st.session_state.get("last_transcription"):
            st.info("Nenhuma transcrição concluída nesta sessão.")
        else:
            final_text = st.session_state["last_transcription"]
            plain_text = st.session_state.get("last_plain_text", final_text)
            lines = final_text.splitlines()
            metadata = st.session_state.get("last_metadata", {})
            segments_data = st.session_state.get("last_segments", [])
            base_name = safe_filename(metadata.get("filename", "transcricao"))

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Arquivo", metadata.get("filename", "-")[:24])
            col2.metric("Duração", metadata.get("duration", "-"))
            col3.metric("Modelo", metadata.get("model", "-"))
            col4.metric("Trechos", str(len(segments_data)))

            st.subheader("Transcrição")
            st.text_area("Resultado", final_text, height=440)

            st.subheader("Downloads")
            dl1, dl2, dl3, dl4 = st.columns(4)
            with dl1:
                st.download_button(
                    "TXT",
                    final_text.encode("utf-8"),
                    file_name=f"{base_name}_transcricao.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
            with dl2:
                docx_bytes = save_docx(lines, metadata)
                st.download_button(
                    "Word",
                    docx_bytes,
                    file_name=f"{base_name}_transcricao.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl3:
                pdf_bytes = save_pdf(lines, metadata)
                st.download_button(
                    "PDF",
                    pdf_bytes,
                    file_name=f"{base_name}_transcricao.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            with dl4:
                srt_text = build_srt(segments_data)
                st.download_button(
                    "SRT",
                    srt_text.encode("utf-8"),
                    file_name=f"{base_name}_legenda.srt",
                    mime="text/plain",
                    use_container_width=True,
                )

    elif page == "Prompts":
        if not st.session_state.get("last_plain_text"):
            st.info("Conclua uma transcrição para gerar os prompts.")
        else:
            plain_text = st.session_state["last_plain_text"]
            st.write("Copie o prompt desejado e cole no ChatGPT para transformar a transcrição.")
            prompts = build_prompts(plain_text)
            for title, prompt in prompts.items():
                with st.expander(title):
                    st.text_area(title, prompt, height=220)
                    st.download_button(
                        f"Baixar prompt - {title}",
                        prompt.encode("utf-8"),
                        file_name=f"prompt_{safe_filename(title)}.txt",
                        mime="text/plain",
                    )


    elif page == "Ferramentas":
        st.subheader("Ferramentas")

        st.error(
            "A fragmentação de arquivos grandes dentro do Streamlit Cloud foi desativada para evitar queda do app. "
            "O upload de arquivos muito grandes consome memória e pode derrubar o sistema no final do processamento."
        )

        st.markdown("### Fragmentador Local")
        st.write(
            "Use este assistente para dividir vídeos, áudios ou arquivos grandes diretamente no seu computador. "
            "Depois envie as partes menores pela aba **Transcrever**."
        )

        c1, c2 = st.columns([1, 1])
        with c1:
            st.download_button(
                "Baixar Fragmentador Local robusto",
                data=build_local_fragmenter_bat(),
                file_name="TranscreveFacil_Fragmentador_Local.bat",
                mime="application/octet-stream",
                use_container_width=True,
            )
        with c2:
            st.info("Recomendação: partes de 10 a 15 minutos para vídeos de reunião, aula ou webinário.")

        st.markdown("### Como usar")
        st.write(
            "1. Baixe o Fragmentador Local.\n\n"
            "2. Dê dois cliques no arquivo `.bat`. Se o FFmpeg não estiver instalado, o assistente tentará baixar uma versão portátil automaticamente.\n\n"
            "3. Arraste o vídeo/áudio para a janela que abrir.\n\n"
            "4. Informe a duração de cada parte, por exemplo, `10` minutos.\n\n"
            "5. As partes serão salvas em `Downloads > TranscreveFacil_Fragmentado`.\n\n"
            "6. Volte ao Transcreve Fácil e envie cada parte pela aba **Transcrever**."
        )

        st.markdown("### Ferramentas em nuvem ainda disponíveis")
        ferramenta = st.radio(
            "Escolha uma ferramenta leve",
            ["Compactar arquivos em ZIP", "Compactar áudio/vídeo pequeno"],
            horizontal=False,
        )

        if ferramenta == "Compactar arquivos em ZIP":
            st.info("Use apenas para arquivos pequenos ou médios. Para arquivos grandes, compacte localmente.")
            zip_files = st.file_uploader(
                "Escolha arquivos para compactar",
                accept_multiple_files=True,
                key="zip_files_safe",
            )
            if zip_files and st.button("Gerar ZIP", type="primary", use_container_width=True):
                total_mb = sum(f.size for f in zip_files) / (1024 * 1024)
                if total_mb > 120:
                    st.error("Conjunto acima de 120 MB. Para evitar erro de rede no Streamlit Cloud, use ferramenta local.")
                    st.stop()
                with tempfile.TemporaryDirectory() as tmpdir:
                    paths = []
                    for f in zip_files:
                        p = os.path.join(tmpdir, safe_filename(f.name) + Path(f.name).suffix)
                        with open(p, "wb") as out:
                            out.write(f.getbuffer())
                        paths.append(p)
                    zip_path = os.path.join(tmpdir, "arquivos_compactados.zip")
                    make_zip_from_paths(paths, zip_path)
                    st.success("ZIP gerado.")
                    st.download_button(
                        "Baixar ZIP",
                        bytes_from_file(zip_path),
                        file_name="arquivos_compactados.zip",
                        mime="application/zip",
                        use_container_width=True,
                    )

        elif ferramenta == "Compactar áudio/vídeo pequeno":
            st.info("Use apenas com arquivos pequenos. Para vídeos grandes, use o Fragmentador Local.")
            uploaded = st.file_uploader(
                "Escolha áudio ou vídeo pequeno",
                type=[ext.replace(".", "") for ext in SUPPORTED_EXTS],
                key="compress_media_safe",
            )
            output_ext = st.selectbox("Formato de saída", [".mp3", ".mp4"], index=0)
            if uploaded and st.button("Compactar mídia", type="primary", use_container_width=True):
                if uploaded.size > 120 * 1024 * 1024:
                    st.error("Arquivo acima de 120 MB. Para evitar erro de rede no Streamlit Cloud, use o Fragmentador Local robusto.")
                    st.stop()
                with tempfile.TemporaryDirectory() as tmpdir:
                    suffix = Path(uploaded.name).suffix.lower()
                    input_path = os.path.join(tmpdir, safe_filename(uploaded.name) + suffix)
                    with open(input_path, "wb") as f:
                        f.write(uploaded.getbuffer())
                    output_path = os.path.join(tmpdir, "midia_compactada" + output_ext)
                    try:
                        compress_media(input_path, output_path, output_ext)
                        original = os.path.getsize(input_path)
                        final = os.path.getsize(output_path)
                        reduction = 100 - (final / original * 100) if original else 0
                        st.success(f"Compactação concluída. Redução aproximada: {reduction:.1f}%.")
                        st.download_button(
                            "Baixar mídia compactada",
                            bytes_from_file(output_path),
                            file_name=os.path.basename(output_path),
                            mime="video/mp4" if output_ext == ".mp4" else "audio/mpeg",
                            use_container_width=True,
                        )
                    except Exception as e:
                        st.error("Não foi possível compactar a mídia.")
                        st.warning(str(e))


    elif page == "Conversao privada":
        render_conversao_privada_page()

    elif page == "YouTube local":
        st.subheader("Modo recomendado para YouTube")
        st.info(
            "Quando o YouTube bloquear o download no Streamlit Cloud, baixe o áudio no seu computador "
            "e depois envie o arquivo pela aba Transcrever > Enviar arquivo. Esse fluxo é mais estável."
        )

        st.markdown("### Passo a passo no Windows")
        st.write("1. Instale o yt-dlp no seu computador, uma única vez:")
        st.code("python -m pip install -U yt-dlp", language="bash")
        st.write("2. Instale o FFmpeg, se ainda não tiver:")
        st.code("winget install Gyan.FFmpeg", language="bash")
        st.write("3. Baixe apenas o áudio do vídeo autorizado:")
        st.code('yt-dlp -x --audio-format mp3 --audio-quality 0 "COLE_A_URL_DO_YOUTUBE_AQUI"', language="bash")
        st.write("4. Depois envie o MP3 gerado pela opção Enviar arquivo.")

        st.markdown("### Gerador de comando")
        local_url = st.text_input("URL para gerar comando local", placeholder="https://www.youtube.com/watch?v=...")
        if local_url:
            st.code(f'yt-dlp -x --audio-format mp3 --audio-quality 0 "{local_url.strip()}"', language="bash")
            st.caption("Cole esse comando no Prompt/Terminal do seu computador, não no Streamlit Cloud.")

        st.markdown("### Por que esse modo é melhor?")
        st.write(
            "O Streamlit Cloud roda em servidor de nuvem. O YouTube pode bloquear esse tipo de acesso com 403 ou pedido de login. "
            "No seu computador, o acesso costuma ser mais estável porque vem da sua própria rede."
        )

    elif page == "Ajuda":
        st.subheader("Como usar bem no Streamlit Cloud")
        st.write(
            "1. Comece com arquivos curtos.\n\n"
            "2. Use o modelo small para maior estabilidade.\n\n"
            "3. Para vídeos longos, prefira rodar a versão local no computador.\n\n"
            "4. O resultado não fica salvo permanentemente. Baixe TXT, Word ou PDF assim que terminar.\n\n"
            "5. Use a aba Ferramentas para fragmentar arquivos longos ou compactar arquivos antes do upload."
        )
        st.subheader("Configuração de usuários em Secrets")
        st.code(
            '[users]\n'
            '"vmsoares@tre-ba.jus.br" = "SUA_SENHA_FORTE"\n\n'
            '[profiles]\n'
            '"vmsoares@tre-ba.jus.br" = "admin"',
            language="toml",
        )


if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    login_screen()
else:
    app_screen()
